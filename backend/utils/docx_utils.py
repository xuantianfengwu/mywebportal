#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docxtpl import DocxTemplate, InlineImage
from ..utils.osfile_utils import OSFileUtils
import docx
from zhconv import convert
from win32com.client import Dispatch, constants, gencache
import pythoncom
import os

class DocxUtils(object):
    """ 操作Word Docx的类 """
    def __init__(self, docx_path=None):
        """ 初始化 """
        self.docx_path = docx_path

    def translate_to_traditional(self, input_path, output_path):
        """ 将Word中所有文字翻译成繁体 """
        doc = docx.Document(input_path)
        # 文章段落
        for parag in doc.paragraphs:
            for run in parag.runs:
                if run.text:
                    if len(run.text.strip()) > 0:
                        ori_text = run.text
                        run.clear()
                        run.add_text(convert(ori_text, 'zh-hk'))
        # 表格
        for table in doc.tables:
            for cell in table._cells:
                for parag in cell.paragraphs:
                    for run in parag.runs:
                        if run.text:
                            if len(run.text.strip()) > 0:
                                ori_text = run.text
                                run.clear()
                                run.add_text(convert(ori_text, 'zh-hk'))
        doc.save(output_path)
        return True

    def docx_to_pdf(self, input_path, output_path):
        """ 将Word转化为PDF """
        # gencache.EnsureModule('{00020905-0000-0000-C000-000000000046}', 0, 8, 4)
        pythoncom.CoInitialize()
        #w = Dispatch('Word.Application')
        w = gencache.EnsureDispatch('Word.Application')
        try:
            # 打开文件
            doc = w.Documents.Open(os.path.abspath(input_path), ReadOnly=1)
            # 转换文件
            doc.ExportAsFixedFormat(os.path.abspath(output_path), constants.wdExportFormatPDF,
                                    Item=constants.wdExportDocumentWithMarkup,
                                    CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
            return True
        except Exception as e:
            print('e:', e)
            print(constants.items())
        finally:
            w.Quit(constants.wdDoNotSaveChanges)

    def replace_docx_tpl(self, input_path, output_path, context):
        """ 用docx-tpl包对word进行渲染 """
        return ''
